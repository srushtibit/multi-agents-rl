"""
DOCX document processor for the unified knowledge base.
Handles Microsoft Word documents with table and formatting extraction.
"""

import os
from typing import Dict, Any, List
import logging
from docx import Document
from docx.shared import Inches
from .base_processor import BaseDocumentProcessor

logger = logging.getLogger(__name__)

class DOCXProcessor(BaseDocumentProcessor):
    """Processor for DOCX files."""
    
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.supported_extensions = ['.docx']
        
        # DOCX-specific configuration
        self.include_tables = True
        self.include_headers_footers = True
        self.preserve_formatting_info = True
    
    def can_process(self, file_path: str) -> bool:
        """Check if file is a DOCX."""
        return file_path.lower().endswith('.docx')
    
    def extract_text(self, file_path: str) -> str:
        """Extract text content from DOCX file."""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Add document title if available
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                text_parts.append(f"Title: {doc.core_properties.title}")
                text_parts.append("")
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Include style information if preserving formatting
                    if self.preserve_formatting_info and paragraph.style.name != 'Normal':
                        text_parts.append(f"[{paragraph.style.name}] {paragraph.text}")
                    else:
                        text_parts.append(paragraph.text)
            
            # Extract text from tables if enabled
            if self.include_tables:
                table_texts = self._extract_table_text(doc)
                if table_texts:
                    text_parts.append("")
                    text_parts.append("=== TABLES ===")
                    text_parts.extend(table_texts)
            
            # Extract headers and footers if enabled
            if self.include_headers_footers:
                header_footer_text = self._extract_headers_footers(doc)
                if header_footer_text:
                    text_parts.append("")
                    text_parts.append("=== HEADERS/FOOTERS ===")
                    text_parts.extend(header_footer_text)
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX file."""
        metadata = {
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'file_size': os.path.getsize(file_path),
            'format': 'docx'
        }
        
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            # Extract core properties
            metadata.update({
                'title': getattr(core_props, 'title', ''),
                'author': getattr(core_props, 'author', ''),
                'subject': getattr(core_props, 'subject', ''),
                'keywords': getattr(core_props, 'keywords', ''),
                'comments': getattr(core_props, 'comments', ''),
                'created': getattr(core_props, 'created', None),
                'modified': getattr(core_props, 'modified', None),
                'last_modified_by': getattr(core_props, 'last_modified_by', ''),
                'revision': getattr(core_props, 'revision', 0),
                'category': getattr(core_props, 'category', ''),
                'language': getattr(core_props, 'language', ''),
            })
            
            # Count document elements
            num_paragraphs = len(doc.paragraphs)
            num_tables = len(doc.tables)
            num_sections = len(doc.sections)
            
            # Count words and characters (approximate)
            total_text = self.extract_text(file_path)
            word_count = len(total_text.split())
            char_count = len(total_text)
            
            metadata.update({
                'num_paragraphs': num_paragraphs,
                'num_tables': num_tables,
                'num_sections': num_sections,
                'word_count': word_count,
                'char_count': char_count,
            })
            
            # Extract style information
            styles_used = set()
            for paragraph in doc.paragraphs:
                if paragraph.style.name:
                    styles_used.add(paragraph.style.name)
            
            metadata['styles_used'] = list(styles_used)
            
            # Extract table information
            if self.include_tables and num_tables > 0:
                table_info = []
                for i, table in enumerate(doc.tables):
                    table_info.append({
                        'table_index': i,
                        'rows': len(table.rows),
                        'columns': len(table.columns) if table.rows else 0
                    })
                metadata['tables'] = table_info
            
        except Exception as e:
            logger.warning(f"Could not extract detailed metadata from {file_path}: {e}")
            metadata['error'] = str(e)
        
        return metadata
    
    def _extract_table_text(self, doc: Document) -> List[str]:
        """Extract text from all tables in the document."""
        table_texts = []
        
        for table_index, table in enumerate(doc.tables):
            try:
                table_text = [f"Table {table_index + 1}:"]
                
                # Extract table headers (first row)
                if table.rows:
                    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                    if any(header_cells):  # Only add if not all empty
                        table_text.append(f"Headers: {' | '.join(header_cells)}")
                
                # Extract table data
                for row_index, row in enumerate(table.rows[1:], 1):  # Skip header row
                    row_cells = [cell.text.strip() for cell in row.cells]
                    if any(row_cells):  # Only add if not all empty
                        table_text.append(f"Row {row_index}: {' | '.join(row_cells)}")
                
                if len(table_text) > 1:  # Only add if table has content
                    table_texts.extend(table_text)
                    table_texts.append("")  # Empty line after each table
                    
            except Exception as e:
                logger.warning(f"Error extracting table {table_index}: {e}")
                continue
        
        return table_texts
    
    def _extract_headers_footers(self, doc: Document) -> List[str]:
        """Extract text from headers and footers."""
        header_footer_texts = []
        
        try:
            for section_index, section in enumerate(doc.sections):
                # Extract header
                if section.header:
                    header_paragraphs = section.header.paragraphs
                    header_text = []
                    for para in header_paragraphs:
                        if para.text.strip():
                            header_text.append(para.text.strip())
                    
                    if header_text:
                        header_footer_texts.append(f"Section {section_index + 1} Header:")
                        header_footer_texts.extend(header_text)
                        header_footer_texts.append("")
                
                # Extract footer
                if section.footer:
                    footer_paragraphs = section.footer.paragraphs
                    footer_text = []
                    for para in footer_paragraphs:
                        if para.text.strip():
                            footer_text.append(para.text.strip())
                    
                    if footer_text:
                        header_footer_texts.append(f"Section {section_index + 1} Footer:")
                        header_footer_texts.extend(footer_text)
                        header_footer_texts.append("")
        
        except Exception as e:
            logger.warning(f"Error extracting headers/footers: {e}")
        
        return header_footer_texts
    
    def _get_paragraph_formatting_info(self, paragraph) -> Dict[str, Any]:
        """Extract formatting information from a paragraph."""
        formatting_info = {}
        
        try:
            # Style information
            formatting_info['style'] = paragraph.style.name
            
            # Alignment
            if paragraph.alignment is not None:
                alignment_map = {
                    0: 'left',
                    1: 'center', 
                    2: 'right',
                    3: 'justify'
                }
                formatting_info['alignment'] = alignment_map.get(paragraph.alignment, 'unknown')
            
            # Check for runs (character-level formatting)
            has_bold = any(run.bold for run in paragraph.runs if run.bold)
            has_italic = any(run.italic for run in paragraph.runs if run.italic)
            has_underline = any(run.underline for run in paragraph.runs if run.underline)
            
            if has_bold:
                formatting_info['has_bold'] = True
            if has_italic:
                formatting_info['has_italic'] = True
            if has_underline:
                formatting_info['has_underline'] = True
            
        except Exception as e:
            logger.debug(f"Could not extract formatting info: {e}")
        
        return formatting_info